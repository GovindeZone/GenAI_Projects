from docx import Document


def export_docx(report, filename):

    document = Document()

    document.add_heading(
        "Business Report",
        level=1
    )

    document.add_paragraph(report)

    document.save(filename)